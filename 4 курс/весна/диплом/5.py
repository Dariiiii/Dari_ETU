import re
import nltk
import docx
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
from collections import Counter

# Загрузка необходимых ресурсов
nltk.download('stopwords')
nltk.download('punkt')
nltk.download('averaged_perceptron_tagger')


class DocxUploader:
    def __init__(self):
        self.file = None
        self.text = ""
        self.paragraphs = []

    def upload(self, file):
        self.file = docx.Document(file)
        self.text = self.extract_text()

    def extract_text(self):
        text = []
        for paragraph in self.file.paragraphs:
            if paragraph.text.strip():  # Убираем пустые параграфы
                text.append(paragraph.text)
        return "\n".join(text)

    def analyze_water(self):
        text = self.text.lower()
        stop_words = set(stopwords.words('russian'))
        sentences = sent_tokenize(text)

        water_sentences_count = 0
        total_sentences = len(sentences)

        for sentence in sentences:
            words = word_tokenize(sentence)
            words = [word for word in words if word.isalnum()]
            stop_words_count = len([word for word in words if word in stop_words])
            total_words = len(words)

            if total_words > 0 and (stop_words_count / total_words) * 100 > 30:
                water_sentences_count += 1

        water_ratio = (water_sentences_count / total_sentences) * 100 if total_sentences > 0 else 0

        return {
            "total_sentences": total_sentences,
            "water_sentences_count": water_sentences_count,
            "water_ratio": water_ratio
        }

    def calculate_sentence_length(self, sentences, length_threshold=20):
        return sum(1 for sentence in sentences if len(word_tokenize(sentence)) > length_threshold)

    def count_repeated_phrases(self, sentences, repeat_threshold=5):
        word_count = Counter()
        for sentence in sentences:
            word_count.update(word_tokenize(sentence))

        return sum(1 for count in word_count.values() if count >= repeat_threshold)

    def count_noun_adjective_combinations(self, sentences):
        noun_adjective_count = 0
        for sentence in sentences:
            words = word_tokenize(sentence)
            tagged = nltk.pos_tag(words, lang='rus')
            noun_adjective_count += sum(1 for i in range(len(tagged) - 1) 
                                         if 'NN' in tagged[i][1] and 'JJ' in tagged[i+1][1])
        return noun_adjective_count

    def count_verbs(self, sentences):
        verb_count = 0
        for sentence in sentences:
            words = word_tokenize(sentence)
            tagged = nltk.pos_tag(words, lang='rus')
            verb_count += sum(1 for word, tag in tagged if 'VB' in tag)
        return verb_count

    def calculate_complex_constructions(self, sentences):
        complex_constructions = 0
        for sentence in sentences:
            words = word_tokenize(sentence)
            tagged = nltk.pos_tag(words, lang='rus')
            complex_constructions += sum(1 for word, tag in tagged if 'V' in tag)
        return complex_constructions

    def print_water_analysis(self):
        sentences = self.text.split('. ')
        result = self.analyze_water()

        long_sentences = self.calculate_sentence_length(sentences)
        repeated_phrases = self.count_repeated_phrases(sentences)
        noun_adjective_combinations = self.count_noun_adjective_combinations(sentences)
        verb_count = self.count_verbs(sentences)
        complex_constructions = self.calculate_complex_constructions(sentences)

        print(f"Общее количество предложений: {result['total_sentences']}")
        print(f"Количество предложений с 'водой': {result['water_sentences_count']}")
        print(f"Процент 'воды' по предложениям: {result['water_ratio']:.2f}%")
        print(f"Количество предложений длиной более 20 слов: {long_sentences}")
        print(f"Количество повторяющихся фраз: {repeated_phrases}")
        print(f"Количество сочетаний существительных и прилагательных: {noun_adjective_combinations}")
        print(f"Количество глаголов в тексте: {verb_count}")
        print(f"Количество сложных конструкций (причастия, деепричастия): {complex_constructions}")

# Пример использования
uploader = DocxUploader()
uploader.upload("2024ВКР030438МАКСИМОВ.docx")
uploader.print_water_analysis()
